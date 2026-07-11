from __future__ import annotations

import logging
import uuid
from datetime import datetime
from pathlib import Path

from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader

from app.core.config import (
    MAX_UPLOAD_SIZE,
    SUPPORTED_EXTENSIONS,
    SUPPORTED_MIME_TYPES,
    UPLOAD_DIR,
)
from app.core.exceptions import (
    DocumentError,
    EmptyDocument,
    FileTooLarge,
    UnsupportedFileType,
)
from app.models.document import DocumentMetadata

logger = logging.getLogger(__name__)


class DocumentLoader:
    """
    Handles document persistence and text extraction.

    Responsibilities:
    - Save uploaded files
    - Extract text from existing files
    """

    @staticmethod
    async def save(
        file: UploadFile,
    ) -> tuple[DocumentMetadata, Path]:
        """
        Save an uploaded document once and return its metadata together
        with the saved file path.
        """

        extension = Path(file.filename).suffix.lower()

        if extension not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFileType(
                f"Unsupported extension: {extension}"
            )

        if file.content_type not in SUPPORTED_MIME_TYPES:
            raise UnsupportedFileType(
                f"Unsupported MIME type: {file.content_type}"
            )

        contents = await file.read()

        if len(contents) > MAX_UPLOAD_SIZE:
            raise FileTooLarge(
                "Maximum upload size exceeded."
            )

        document_id = str(uuid.uuid4())

        filename = f"{document_id}{extension}"

        destination = UPLOAD_DIR / filename

        destination.write_bytes(contents)

        metadata = DocumentMetadata(
            document_id=document_id,
            filename=file.filename,
            stored_filename=filename,
            extension=extension,
            size_bytes=len(contents),
            character_count=0,
            word_count=0,
            chunk_count=0,
            uploaded_at=datetime.utcnow(),
            status="uploaded",
        )

        logger.info(
            "Saved document %s (%s)",
            metadata.document_id,
            metadata.filename,
        )

        return metadata, destination

    @staticmethod
    def extract_text(
        path: Path,
    ) -> str:
        """
        Extract readable text from an already-saved document.
        """

        if path.suffix == ".pdf":
            text = DocumentLoader._read_pdf(path)

        elif path.suffix == ".docx":
            text = DocumentLoader._read_docx(path)

        elif path.suffix == ".txt":
            text = DocumentLoader._read_txt(path)

        else:
            raise UnsupportedFileType(
                f"Unsupported extension: {path.suffix}"
            )

        if not text.strip():
            raise EmptyDocument(
                "Document contains no readable text."
            )

        return text

    @staticmethod
    async def process(
        file: UploadFile,
    ) -> tuple[DocumentMetadata, str]:
        """
        Backward compatibility wrapper.

        New code should use:
            save()
            extract_text()
        """

        metadata, saved_path = await DocumentLoader.save(file)

        text = DocumentLoader.extract_text(saved_path)

        metadata.character_count = len(text)
        metadata.word_count = len(text.split())
        metadata.status = "processed"

        logger.info(
            "Processed document %s",
            metadata.document_id,
        )

        return metadata, text

    @staticmethod
    def _read_pdf(
        path: Path,
    ) -> str:

        reader = PdfReader(path)

        pages: list[str] = []

        for page in reader.pages:

            text = page.extract_text()

            if text:
                pages.append(text)

        return "\n".join(pages)

    @staticmethod
    def _read_docx(
        path: Path,
    ) -> str:

        doc = Document(path)

        return "\n".join(
            paragraph.text
            for paragraph in doc.paragraphs
        )

    @staticmethod
    def _read_txt(
        path: Path,
    ) -> str:

        return path.read_text(
            encoding="utf-8",
            errors="ignore",
        )